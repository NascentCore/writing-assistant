import PyPDF2
from docx import Document
from docx.document import Document as _Document
from docx.oxml.text.paragraph import CT_P
from docx.text.paragraph import Paragraph
from docx.oxml.table import CT_Tbl
from docx.table import _Cell, Table, _Row
from typing import List, Union, Optional, Dict, Any, Tuple
import os
import re

class DocumentParser:
    """文档解析器基类"""
    def parse(self, file_path: str) -> str:
        raise NotImplementedError

    def parse_to_doc(self, file_path: str) -> Union[Document, str]:
        raise NotImplementedError
    
class PDFEncryptedError(Exception):
    """当 PDF 文件加密且无法解析时抛出的异常"""
    pass

class PDFParser(DocumentParser):
    """PDF文档解析器"""
    def parse(self, file_path: str) -> str:
        try:
            # 创建PDF reader对象
            pdf_reader = PyPDF2.PdfReader(file_path)
            
            # 检查是否加密
            if pdf_reader.is_encrypted:
                try:
                    # 尝试用空密码解密（对于只有权限密码的PDF通常可行）
                    pdf_reader.decrypt('')
                    
                    # 尝试提取文本
                    text_content = []
                    for page in pdf_reader.pages:
                        text_content.append(page.extract_text())
                    return "\n".join(text_content)
                except:
                    raise PDFEncryptedError("无法解析加密的PDF文件，请先解除文件加密后再上传")
                
            # 提取所有页面的文本
            text_content = []
            for page in pdf_reader.pages:
                text_content.append(page.extract_text())
                
            return "\n".join(text_content)
        except Exception as e:
            if "PyCryptodome is required for AES algorithm" in str(e):
                raise PDFEncryptedError("无法解析加密的PDF文件，请先解除文件加密后再上传")
            elif "File has not been decrypted" in str(e):
                raise PDFEncryptedError("无法解析加密的PDF文件，请先解除文件加密后再上传")
            else:
                raise Exception(f"PDF解析错误: {str(e)}")

class DocxParser(DocumentParser):
    """Word文档解析器"""
    def parse(self, file_path: str) -> str:
        try:
            # 加载Word文档
            doc = Document(file_path)
            
            # 提取所有段落的文本
            text_content = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():  # 只添加非空段落
                    text_content.append(paragraph.text)
                    
            return "\n".join(text_content)
        except Exception as e:
            raise Exception(f"Word文档解析错误: {str(e)}")

    def parse_to_doc(self, file_path: str) -> Document:
        """
        解析Word文档，返回Document对象
        
        Args:
            file_path: 文件路径
            
        Returns:
            Document: python-docx的Document对象
        """
        try:
            doc = Document(file_path)
            # 保存文件路径，以便后续使用
            doc._path = file_path
            return doc
        except Exception as e:
            raise Exception(f"Word文档解析错误: {str(e)}")
    
    def get_outline_structure(self, doc: Document) -> Dict[str, Any]:
        """
        从Word文档中提取大纲结构
        
        Args:
            doc: python-docx的Document对象
            
        Returns:
            Dict: 包含大纲结构的字典
        """
        # 提取标题（使用文件名，移除后缀）
        title = "未命名大纲"
        if hasattr(doc, 'core_properties') and doc.core_properties.title:
            title = doc.core_properties.title
        elif hasattr(doc, '_path') and doc._path:
            # 从文件路径中提取文件名并移除后缀
            file_name = os.path.basename(doc._path)
            title = os.path.splitext(file_name)[0]
        
        # 构建大纲结构
        outline_data = {
            "title": title,
            "sub_paragraphs": []
        }
        
        # 遍历段落，根据段落样式和缩进级别构建大纲结构
        paragraphs_stack = []
        
        for paragraph in doc.paragraphs:  # 不再跳过第一个段落，因为现在使用文件名作为标题
            if not paragraph.text.strip():
                continue
            
            # 获取段落级别
            level = self._get_paragraph_level(paragraph)
            
            # 构建段落数据
            para_data = {
                "title": paragraph.text,
                "description": self._get_paragraph_description(paragraph),
                "level": level,
                "children": []
            }
            
            # 如果是一级段落，添加count_style
            if level == 1:
                para_data["count_style"] = "medium"
            
            # 根据层级关系添加到正确的父节点
            while paragraphs_stack and paragraphs_stack[-1]["level"] >= level:
                paragraphs_stack.pop()
            
            if paragraphs_stack:
                if "children" not in paragraphs_stack[-1]:
                    paragraphs_stack[-1]["children"] = []
                paragraphs_stack[-1]["children"].append(para_data)
            else:
                outline_data["sub_paragraphs"].append(para_data)
            
            paragraphs_stack.append(para_data)
        
        return outline_data
    
    def _get_paragraph_level(self, paragraph: Paragraph) -> int:
        """
        获取段落的层级
        
        Args:
            paragraph: 段落对象
            
        Returns:
            int: 段落层级（1-6）
        """
        try:
            # 首先检查是否是标题样式
            if hasattr(paragraph, 'style') and paragraph.style and paragraph.style.name:
                if paragraph.style.name.startswith('Heading'):
                    try:
                        level = int(paragraph.style.name[-1])
                        if 1 <= level <= 6:
                            return level
                    except (ValueError, IndexError):
                        pass
            
            # 检查大纲级别
            try:
                if (paragraph._element is not None and 
                    hasattr(paragraph._element, 'pPr') and 
                    paragraph._element.pPr is not None):
                    pPr = paragraph._element.pPr
                    if hasattr(pPr, 'outlineLvl') and pPr.outlineLvl is not None:
                        val = pPr.outlineLvl.val
                        if val is not None:
                            return min(int(val) + 1, 6)
            except (AttributeError, ValueError):
                pass
            
            # 根据缩进判断级别
            try:
                if (hasattr(paragraph, 'paragraph_format') and 
                    paragraph.paragraph_format is not None and
                    paragraph.paragraph_format.left_indent is not None):
                    indent = paragraph.paragraph_format.left_indent.pt
                    if indent > 0:
                        return min(max(int(indent / 36) + 1, 1), 6)  # 36pt = 0.5英寸
            except (AttributeError, ValueError):
                pass
            
            # 检查是否使用项目符号或编号
            try:
                if (hasattr(paragraph, '_p') and 
                    paragraph._p is not None and 
                    hasattr(paragraph._p, 'pPr') and 
                    paragraph._p.pPr is not None and
                    hasattr(paragraph._p.pPr, 'numPr') and 
                    paragraph._p.pPr.numPr is not None and
                    hasattr(paragraph._p.pPr.numPr, 'ilvl') and 
                    paragraph._p.pPr.numPr.ilvl is not None and
                    hasattr(paragraph._p.pPr.numPr.ilvl, 'val')):
                    val = paragraph._p.pPr.numPr.ilvl.val
                    if val is not None:
                        return min(int(val) + 1, 6)
            except (AttributeError, ValueError):
                pass
            
            return 1
            
        except Exception as e:
            print(f"获取段落级别时出错: {str(e)}")  # 记录错误但不影响处理
            return 1  # 出错时返回默认级别
    
    def _get_paragraph_description(self, paragraph: Paragraph) -> str:
        """
        获取段落的描述（查找下一个非空段落直到遇到新的标题）
        
        Args:
            paragraph: 段落对象
            
        Returns:
            str: 段落描述
        """
        description_parts = []
        next_p = paragraph._element.getnext()
        
        while (next_p is not None and 
               isinstance(next_p, CT_P) and 
               not Paragraph(next_p, paragraph._parent).style.name.startswith('Heading')):
            next_para = Paragraph(next_p, paragraph._parent)
            if next_para.text.strip():
                description_parts.append(next_para.text)
            next_p = next_p.getnext()
        
        return "\n".join(description_parts)

class MarkdownParser(DocumentParser):
    """Markdown文档解析器"""
    
    def parse(self, file_path: str) -> str:
        """
        解析Markdown文件并返回文本内容
        
        Args:
            file_path: Markdown文件路径
            
        Returns:
            str: 文件内容
        """
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except Exception as e:
            raise Exception(f"Markdown解析错误: {str(e)}")
    
    def parse_to_doc(self, file_path: str) -> Dict[str, Any]:
        """
        解析Markdown文件并返回文本内容和文件路径
        
        Args:
            file_path: Markdown文件路径
            
        Returns:
            Dict: 包含文件内容和文件路径的字典
        """
        content = self.parse(file_path)
        return {
            "content": content,
            "file_path": file_path
        }
    
    def get_outline_structure(self, doc: Dict[str, Any]) -> Dict[str, Any]:
        """
        从Markdown内容中提取大纲结构
        
        Args:
            doc: 包含Markdown内容和文件路径的字典
            
        Returns:
            Dict: 包含大纲结构的字典
        """
        content = doc["content"]
        file_path = doc["file_path"]
        
        # 提取标题（使用文件名，移除后缀）
        title = "未命名大纲"
        if file_path:
            # 从文件路径中提取文件名并移除后缀
            file_name = os.path.basename(file_path)
            title = os.path.splitext(file_name)[0]
        
        # 构建大纲结构
        outline_data = {
            "title": title,
            "sub_paragraphs": []
        }
        
        # 解析Markdown内容
        lines = content.split('\n')
        
        # 提取标题和内容
        current_heading = None
        current_level = 0
        current_content = []
        
        # 存储段落的堆栈，用于构建层级关系
        # 每个元素是 (level, paragraph_data) 的元组
        paragraphs_stack = []
        
        for line in lines:
            # 检查是否是标题行
            heading_match = re.match(r'^(#{1,6})\s+(.+)$', line)
            
            if heading_match:
                # 如果有之前的标题，先处理它
                if current_heading:
                    para_data = {
                        "title": current_heading,
                        "description": "\n".join(current_content).strip(),
                        "level": current_level,
                        "children": []
                    }
                    
                    # 如果是一级段落，添加count_style
                    if current_level == 1:
                        para_data["count_style"] = "medium"
                    
                    # 根据层级关系添加到正确的父节点
                    # 弹出所有级别大于或等于当前级别的段落
                    while paragraphs_stack and paragraphs_stack[-1][0] >= current_level:
                        paragraphs_stack.pop()
                    
                    # 添加到父节点或根节点
                    if paragraphs_stack:
                        # 添加到父节点的children列表
                        paragraphs_stack[-1][1]["children"].append(para_data)
                    else:
                        # 添加到根节点
                        outline_data["sub_paragraphs"].append(para_data)
                    
                    # 将当前段落添加到堆栈
                    paragraphs_stack.append((current_level, para_data))
                
                # 处理新标题
                level = len(heading_match.group(1))
                heading = heading_match.group(2).strip()
                
                current_heading = heading
                current_level = level
                current_content = []
            else:
                # 非标题行，添加到当前内容
                if current_heading is not None:
                    current_content.append(line)
        
        # 处理最后一个标题
        if current_heading:
            para_data = {
                "title": current_heading,
                "description": "\n".join(current_content).strip(),
                "level": current_level,
                "children": []
            }
            
            # 如果是一级段落，添加count_style
            if current_level == 1:
                para_data["count_style"] = "medium"
            
            # 根据层级关系添加到正确的父节点
            while paragraphs_stack and paragraphs_stack[-1][0] >= current_level:
                paragraphs_stack.pop()
            
            if paragraphs_stack:
                paragraphs_stack[-1][1]["children"].append(para_data)
            else:
                outline_data["sub_paragraphs"].append(para_data)
        
        return outline_data

def get_parser(file_type: str) -> DocumentParser:
    """
    根据文件类型获取对应的解析器
    
    Args:
        file_type: 文件类型（扩展名）
        
    Returns:
        DocumentParser: 对应的文档解析器
    """
    file_type = file_type.lower()
    
    if file_type == '.pdf':
        return PDFParser()
    elif file_type in ['.doc', '.docx']:
        return DocxParser()
    elif file_type == '.md':
        return MarkdownParser()
    else:
        raise ValueError(f"不支持的文件类型: {file_type}")

def get_file_format(file_path: str) -> str:
    """
    通过读取文件头部信息来判断文件格式
    
    Args:
        file_path: 文件路径
        
    Returns:
        str: 文件格式 ('pdf' 或 'docx')
    """
    with open(file_path, 'rb') as f:
        header = f.read(4)
        
    # PDF文件头: %PDF (25 50 44 46)
    if header.startswith(b'%PDF'):
        return 'pdf'
    # DOCX文件头: PK (50 4B)
    elif header.startswith(b'PK'):
        return 'docx'
    else:
        raise ValueError(f"不支持的文件格式")
