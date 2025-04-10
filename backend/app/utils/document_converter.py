# -*- coding: utf-8 -*-
from io import BytesIO
from typing import Optional, List
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
import re

def add_numbering_to_headers(html_text, numbering_type):
    # 使用 BeautifulSoup 解析 HTML 内容
    soup = BeautifulSoup(html_text, 'html.parser')
    
    # 中文一、二、三等
    chinese_numbers = ["一", "二", "三", "四", "五", "六", "七", "八", "九", "十",
                       "十一", "十二", "十三", "十四", "十五", "十六", "十七", "十八", "十九", "二十",
                       "二十一", "二十二", "二十三", "二十四", "二十五", "二十六", "二十七", "二十八", "二十九", "三十",
                       "三十一", "三十二", "三十三", "三十四", "三十五", "三十六", "三十七", "三十八", "三十九", "四十",
                       "四十一", "四十二", "四十三", "四十四", "四十五", "四十六", "四十七", "四十八", "四十九", "五十"]
    at_numbers = [
        "①", "②", "③", "④", "⑤", "⑥", "⑦", "⑧", "⑨", "⑩",
        "⑪", "⑫", "⑬", "⑭", "⑮", "⑯", "⑰", "⑱", "⑲", "⑳",
        "㉑", "㉒", "㉓", "㉔", "㉕", "㉖", "㉗", "㉘", "㉙", "㉚",
        "㉛", "㉜", "㉝", "㉞", "㉟", "㊱", "㊲", "㊳", "㊴", "㊵",
        "㊶", "㊷", "㊸", "㊹", "㊺", "㊻", "㊼", "㊽", "㊾", "㊿"
    ]
    
    # 获取中文数字，处理超出范围的情况
    def get_chinese_number(index):
        if index < len(chinese_numbers):
            return chinese_numbers[index]
        else:
            return str(index + 1)
        
    # 获取at数字，处理超出范围的情况
    def get_at_number(index):
        if index < len(at_numbers):
            return at_numbers[index]
        else:
            return str(index + 1)
    
    # 清除所有标题标签的内容中的编号
    for tag in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6']):
        original_text = tag.get_text().strip()
        # 去除现有的编号（中文编号和数字编号）
        clean_text = re.sub(r'^[\d一二三四五六七八九十]+[\.、][\d一二三四五六七八九十\.、]*\s*', '', original_text)
        clean_text = re.sub(r'^\d+(\.\d+)+\s+', '', clean_text)
        tag.clear()
        tag.append(clean_text)  # 保存清理后的文本
    
    # 按文档顺序收集所有标题标签
    all_headers = []
    for tag in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6']):
        level = int(tag.name[1])
        text = tag.get_text().strip()
        all_headers.append((tag, level, text))
    
    # 用于跟踪标题编号的计数器
    h2_counter = 0  # 二级标题计数器
    h3_counters = {}  # 每个二级标题下的三级标题计数器
    h4_counters = {}  # 每个三级标题下的四级标题计数器
    h5_counters = {}  # 每个四级标题下的五级标题计数器
    h6_counters = {}  # 每个五级标题下的六级标题计数器
    
    # 当前标题的父标题ID
    current_h2_id = None
    current_h3_id = None
    current_h4_id = None
    current_h5_id = None
    
    # 为每个标题生成编号
    for i, (tag, level, text) in enumerate(all_headers):
        if level == 1:
            # 一级标题不加编号
            continue
        elif level == 2:
            # 二级标题：一、二、三...
            h2_counter += 1
            current_h2_id = h2_counter
            # 重置低级标题计数器
            h3_counters[current_h2_id] = 0
            
            # 生成编号
            if numbering_type == "chinese":
                numbering = f'第{get_chinese_number(h2_counter-1)}章、'
            elif numbering_type == "mix":
                numbering = f'{get_chinese_number(h2_counter-1)}、'
            else:
                numbering = f'{h2_counter}'
            tag.clear()
            tag.append(f"{numbering} {text}")
        
        elif level == 3:
            # 三级标题：1.1, 1.2, 2.1...
            if current_h2_id is None:
                current_h2_id = 1
                h3_counters[current_h2_id] = 0
            
            h3_counters[current_h2_id] += 1
            current_h3_id = f"{current_h2_id}.{h3_counters[current_h2_id]}"
            # 重置四级标题计数器
            h4_counters[current_h3_id] = 0
            
            # 生成编号
            if numbering_type == "chinese":
                numbering = f'第{get_chinese_number(h3_counters[current_h2_id]-1)}节、'
            elif numbering_type == "mix":
                numbering = f'（{get_chinese_number(h3_counters[current_h2_id]-1)}）'
            else:
                numbering = f'{current_h2_id}.{h3_counters[current_h2_id]}'
            tag.clear()
            tag.append(f"{numbering} {text}")
        
        elif level == 4:
            # 四级标题：1.1.1, 1.1.2, 1.2.1...
            if current_h3_id is None:
                if current_h2_id is None:
                    current_h2_id = 1
                    h3_counters[current_h2_id] = 0
                h3_counters[current_h2_id] += 1
                current_h3_id = f"{current_h2_id}.{h3_counters[current_h2_id]}"
                h4_counters[current_h3_id] = 0
            
            h4_counters[current_h3_id] += 1
            current_h4_id = f"{current_h3_id}.{h4_counters[current_h3_id]}"
            # 重置五级标题计数器
            h5_counters[current_h4_id] = 0
            
            # 生成编号
            if numbering_type == "chinese":
                numbering = f'{get_chinese_number(h4_counters[current_h3_id]-1)}、'
            elif numbering_type == "mix":
                numbering = f'{h4_counters[current_h3_id]}.'
            else:
                numbering = f'{current_h3_id}.{h4_counters[current_h3_id]}'
            tag.clear()
            tag.append(f"{numbering} {text}")
        
        elif level == 5:
            # 五级标题：1.1.1.1, 1.1.1.2...
            if current_h4_id is None:
                if current_h3_id is None:
                    if current_h2_id is None:
                        current_h2_id = 1
                        h3_counters[current_h2_id] = 0
                    h3_counters[current_h2_id] += 1
                    current_h3_id = f"{current_h2_id}.{h3_counters[current_h2_id]}"
                    h4_counters[current_h3_id] = 0
                h4_counters[current_h3_id] += 1
                current_h4_id = f"{current_h3_id}.{h4_counters[current_h3_id]}"
                h5_counters[current_h4_id] = 0
            
            h5_counters[current_h4_id] += 1
            current_h5_id = f"{current_h4_id}.{h5_counters[current_h4_id]}"
            # 重置六级标题计数器
            h6_counters[current_h5_id] = 0
            
            # 生成编号
            if numbering_type == "chinese":
                numbering = f'{h5_counters[current_h4_id]}.'
            elif numbering_type == "mix":
                numbering = f'（{h5_counters[current_h4_id]}）'
            else:
                numbering = f'{current_h4_id}.{h5_counters[current_h4_id]}'
            tag.clear()
            tag.append(f"{numbering} {text}")
            
        elif level == 6:
            # 六级标题：1.1.1.1.1, 1.1.1.1.2...
            if current_h5_id is None:
                if current_h4_id is None:
                    if current_h3_id is None:
                        if current_h2_id is None:
                            current_h2_id = 1
                            h3_counters[current_h2_id] = 0
                        h3_counters[current_h2_id] += 1
                        current_h3_id = f"{current_h2_id}.{h3_counters[current_h2_id]}"
                        h4_counters[current_h3_id] = 0
                    h4_counters[current_h3_id] += 1
                    current_h4_id = f"{current_h3_id}.{h4_counters[current_h3_id]}"
                    h5_counters[current_h4_id] = 0
                h5_counters[current_h4_id] += 1
                current_h5_id = f"{current_h4_id}.{h5_counters[current_h4_id]}"
                h6_counters[current_h5_id] = 0
            
            h6_counters[current_h5_id] += 1
            
            # 生成编号
            if numbering_type == "chinese":
                numbering = f'（{h6_counters[current_h5_id]}）'
            elif numbering_type == "mix":
                numbering = f'{get_at_number(h6_counters[current_h5_id]-1)}'
            else:
                numbering = f'{current_h5_id}.{h6_counters[current_h5_id]}'
            tag.clear()
            tag.append(f"{numbering} {text}")
    
    # 返回处理后的 HTML
    return str(soup)

def fix_document_numbering(doc):
    """
    修复文档中的标题编号问题，确保标题层级正确：
    - 形如"二、"的标题为二级标题
    - 形如"2.3"的标题为三级标题
    - 形如"9.2.1"的标题为四级标题
    
    Args:
        doc: 要修复的文档对象
    
    Returns:
        修复的标题数量
    """
    # 记录修复的标题数量
    fixed_count = 0
    total_paragraphs = len(doc.paragraphs)
    
    # 遍历所有段落
    for i, para in enumerate(doc.paragraphs):
        # 获取标题文本
        text = para.text.strip()
        
        # 检查是否是形如"1.1.1.1.1"的五级标题 (五个数字)
        five_level_match = re.match(r'^(\d+)\.(\d+)\.(\d+)\.(\d+)\.(\d+)[\s\.]*(.*)$', text)
        if five_level_match:
            section, subsection, subsubsection, subsubsubsection, subsubsubsubsection, title_text = five_level_match.groups()
            
            # 清除段落内容
            para.clear()
            
            # 添加标题文本
            run = para.add_run(f"{section}.{subsection}.{subsubsection}.{subsubsubsection}.{subsubsubsubsection} {title_text}")
            
            # 设置字体
            font = run.font
            font.size = Pt(11)
            run.bold = True
            
            # 设置对齐方式
            para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            
            # 使用五级标题样式
            para.style = 'Heading 5'
            
            fixed_count += 1
            continue

        # 检查是否是形如"1.1.1.1"的四级标题 (四个数字)
        four_level_match = re.match(r'^(\d+)\.(\d+)\.(\d+)\.(\d+)[\s\.]*(.*)$', text)
        if four_level_match:
            section, subsection, subsubsection, subsubsubsection, title_text = four_level_match.groups()
            
            # 清除段落内容
            para.clear()
            
            # 添加标题文本
            run = para.add_run(f"{section}.{subsection}.{subsubsection}.{subsubsubsection} {title_text}")
            
            # 设置字体
            font = run.font
            font.size = Pt(12)
            # font.color.rgb = RGBColor(0, 64, 128)
            run.bold = True
            
            # 设置对齐方式
            para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            
            # 使用四级标题样式
            para.style = 'Heading 4'
            
            fixed_count += 1
            continue

        # 检查是否是形如"1.1.1"的三级标题 (三个数字)
        three_level_match = re.match(r'^(\d+)\.(\d+)\.(\d+)[\s\.]*(.*)$', text)
        if three_level_match and not four_level_match:
            section, subsection, subsubsection, title_text = three_level_match.groups()
            
            # 清除段落内容
            para.clear()
            
            # 添加标题文本
            run = para.add_run(f"{section}.{subsection}.{subsubsection} {title_text}")
            
            # 设置字体
            font = run.font
            font.size = Pt(13)
            # font.color.rgb = RGBColor(0, 64, 128)
            run.bold = True
            
            # 设置对齐方式
            para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            
            # 使用三级标题样式
            para.style = 'Heading 3'
            
            fixed_count += 1
            continue
        
        # 检查是否是形如"2.3"或"3.1"的二级标题
        two_level_match = re.match(r'^(\d+)\.(\d+)[\s\.]*(.*)$', text)
        
        # 也检查中文标题格式，如"二.三"
        if not two_level_match:
            chinese_match = re.match(r'^([一二三四五六七八九十]+)[\.．。]([一二三四五六七八九十]+)[\s\.]*(.*)$', text)
            if chinese_match:
                section, subsection, title_text = chinese_match.groups()
                para.clear()
                run = para.add_run(f"{section}.{subsection} {title_text}")
                
                font = run.font
                font.size = Pt(13)
                # font.color.rgb = RGBColor(0, 64, 128)
                run.bold = True
                
                para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                para.style = 'Heading 2'
                
                para.paragraph_format.left_indent = Pt(0)
                para.paragraph_format.first_line_indent = Pt(0)
                
                fixed_count += 1
                continue
        
        if two_level_match:
            section, subsection, title_text = two_level_match.groups()
            para.clear()
            run = para.add_run(f"{section}.{subsection} {title_text}")
            
            font = run.font
            font.size = Pt(13)
            # font.color.rgb = RGBColor(0, 64, 128)
            run.bold = True
            
            para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            para.style = 'Heading 2'
            
            para.paragraph_format.left_indent = Pt(0)
            para.paragraph_format.first_line_indent = Pt(0)
            
            fixed_count += 1
    
    return fixed_count

def process_table(table_element, doc):
    """
    处理HTML表格，转换为Word表格
    
    Args:
        table_element: BeautifulSoup的表格元素
        doc: Word文档对象
    """
    # 创建Word表格
    rows = table_element.find_all('tr')
    if not rows:
        return
        
    # 获取列数（使用第一行的单元格数）
    first_row = rows[0]
    cols = len(first_row.find_all(['td', 'th']))
    
    # 创建表格
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = 'Table Grid'  # 使用网格样式
    
    # 设置表格对齐方式
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # 处理每一行
    for i, row in enumerate(rows):
        cells = row.find_all(['td', 'th'])
        for j, cell in enumerate(cells):
            if j >= cols:  # 跳过超出列数的单元格
                continue
            
            # 设置单元格内容
            word_cell = table.cell(i, j)
            paragraph = word_cell.paragraphs[0]
            
            # 获取单元格样式
            is_header = cell.name == 'th'
            bg_color = cell.get('bgcolor', '')
            align = cell.get('align', 'left')
            
            # 设置单元格对齐方式
            if align == 'center':
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            elif align == 'right':
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            else:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

            # 处理单元格中的内容
            def process_cell_content(element, current_paragraph):
                if isinstance(element, str):
                    text = element.strip()
                    if text:
                        run = current_paragraph.add_run(text)
                        font = run.font
                        font.size = Pt(10.5)
                        if is_header:
                            run.bold = True
                    return
                
                if element.name == 'img':
                    try:
                        src = element.get('src', '')
                        if src:
                            # 获取页面宽度（单位：英寸）
                            section = doc.sections[0]
                            page_width = float(section.page_width.inches)
                            margin_left = float(section.left_margin.inches)
                            margin_right = float(section.right_margin.inches)
                            
                            # 计算可用宽度（页面宽度减去左右页边距）
                            available_width = page_width - margin_left - margin_right
                            
                            # 设置单元格中图片的最大宽度（可用宽度的30%）
                            max_width = available_width * 0.3
                            
                            if src.startswith('data:image'):
                                # 处理base64编码的图片
                                import base64
                                import io
                                header, encoded = src.split(',', 1)
                                image_data = base64.b64decode(encoded)
                                image_stream = io.BytesIO(image_data)
                                current_paragraph.add_run().add_picture(image_stream, width=Inches(max_width))
                            elif src.startswith(('http://', 'https://')):
                                # 处理网络图片
                                import requests
                                response = requests.get(src)
                                image_stream = io.BytesIO(response.content)
                                current_paragraph.add_run().add_picture(image_stream, width=Inches(max_width))
                            else:
                                # 处理本地图片
                                current_paragraph.add_run().add_picture(src, width=Inches(max_width))
                            
                            # 如果图片有标题，添加标题
                            if element.get('alt'):
                                caption_paragraph = word_cell.add_paragraph()
                                caption_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                                caption_run = caption_paragraph.add_run(element.get('alt'))
                                caption_run.italic = True
                                caption_run.font.size = Pt(9)
                    except Exception as e:
                        import logging
                        logging.error(f"处理表格中的图片失败: {str(e)}")
                    return
                
                if element.name == 'p':
                    # 为段落内容创建新的段落
                    new_paragraph = word_cell.add_paragraph()
                    new_paragraph.alignment = paragraph.alignment
                    for child in element.children:
                        process_cell_content(child, new_paragraph)
                    return
                
                if element.name in ['strong', 'b']:
                    run = current_paragraph.add_run(element.get_text().strip())
                    run.bold = True
                    return
                
                if element.name in ['em', 'i']:
                    run = current_paragraph.add_run(element.get_text().strip())
                    run.italic = True
                    return
                
                if element.name == 'br':
                    current_paragraph.add_run('\n')
                    return
                
                # 处理其他元素的子元素
                for child in element.children:
                    process_cell_content(child, current_paragraph)

            # 清除默认段落
            paragraph.clear()
            
            # 处理单元格中的内容
            for content in cell.children:
                process_cell_content(content, paragraph)
            
            # 设置单元格背景色
            if is_header and bg_color:
                try:
                    if bg_color.startswith('#'):
                        color = bg_color[1:]
                        r = int(color[0:2], 16)
                        g = int(color[2:4], 16)
                        b = int(color[4:6], 16)
                    else:
                        color_map = {
                            'gray': (128, 128, 128),
                            'lightgray': (192, 192, 192),
                            'darkgray': (64, 64, 64),
                            'black': (0, 0, 0),
                            'white': (255, 255, 255),
                            'red': (255, 0, 0),
                            'green': (0, 255, 0),
                            'blue': (0, 0, 255),
                            'yellow': (255, 255, 0),
                            'cyan': (0, 255, 255),
                            'magenta': (255, 0, 255)
                        }
                        r, g, b = color_map.get(bg_color.lower(), (255, 255, 255))
                    
                    word_cell._tc.get_or_add_tcPr().append(parse_xml(f'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="{bg_color}"/>'))
                except:
                    pass
            
            # 设置单元格垂直对齐
            word_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

def process_image(img_element, doc):
    """
    处理HTML图片元素，将其添加到Word文档中
    
    Args:
        img_element: BeautifulSoup的图片元素
        doc: Word文档对象
    """
    src = img_element.get('src', '')
    if not src:
        return
    
    try:
        # 获取页面宽度（单位：英寸）
        section = doc.sections[0]
        page_width = float(section.page_width.inches)
        margin_left = float(section.left_margin.inches)
        margin_right = float(section.right_margin.inches)
        
        # 计算可用宽度（页面宽度减去左右页边距）
        available_width = page_width - margin_left - margin_right
        
        # 设置最大宽度为可用宽度的90%
        max_width = available_width * 0.9
        
        # 处理不同类型的图片源
        if src.startswith('data:image'):
            # 处理base64编码的图片
            import base64
            import iof
            # 提取base64数据
            header, encoded = src.split(',', 1)
            image_data = base64.b64decode(encoded)
            image_stream = io.BytesIO(image_data)
            
            # 添加图片到文档并控制大小
            doc.add_picture(image_stream, width=Inches(max_width))
            
        elif src.startswith(('http://', 'https://')):
            # 处理网络图片
            import requests
            response = requests.get(src)
            image_stream = io.BytesIO(response.content)
            
            # 添加图片到文档并控制大小
            doc.add_picture(image_stream, width=Inches(max_width))
            
        else:
            # 处理本地图片
            doc.add_picture(src, width=Inches(max_width))
        
        # 获取最后添加的段落（图片所在段落）
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # 如果图片有标题，添加标题
        if img_element.get('alt'):
            caption = doc.add_paragraph()
            caption.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            caption.add_run(img_element.get('alt')).italic = True
            
    except Exception as e:
        import logging
        logging.error(f"处理图片失败: {str(e)}")

def html_to_docx(
    html_content: str, 
    title: str = "Document", 
    author: str = "System",
    versions: Optional[List[dict]] = None,
    add_numbering: bool = True,
    numbering_type: str = "number",
) -> BytesIO:
    """
    将HTML内容转换为DOCX格式
    
    Args:
        html_content: HTML格式的文档内容
        title: 文档标题
        author: 文档作者
        versions: 可选的版本历史列表，每个版本是一个包含version, content, comment, created_at的字典
        add_numbering: 是否为标题添加序号，默认为True
        
    Returns:
        BytesIO: 包含DOCX文件内容的二进制流
    """
    # 将 HTML 内容转换为 DOCX 文件
    doc = Document()

    # 修改标题样式的默认颜色为黑色
    for i in range(1, 10):  # 处理 Heading 1 到 Heading 9
        style_name = f'Heading {i}'
        if style_name in doc.styles:
            style = doc.styles[style_name]
            style.font.color.rgb = RGBColor(0, 0, 0)  # 设置为黑色
    
    # 预处理HTML，合并列表项后的内容
    def preprocess_html(html_text):
        # 使用正则表达式查找列表项模式并合并后续内容
        pattern = r'(\d+\.)\s*\n+\s*([^<>\d\n][^<>\n]*)'
        processed_html = re.sub(pattern, r'\1 \2', html_text)
        
        # 处理HTML中的有序列表
        soup = BeautifulSoup(processed_html, 'html.parser')
        for ol in soup.find_all('ol'):
            for li in ol.find_all('li'):
                for p in li.find_all('p'):
                    p.replace_with(p.get_text())
        
        return str(soup)
    
    # 预处理HTML
    preprocessed_html = preprocess_html(html_content)
    
    # 检查文档是否有h1标题
    soup_check = BeautifulSoup(preprocessed_html, 'html.parser')
    has_h1_title = bool(soup_check.find('h1'))
    
    # 根据add_numbering参数决定是否添加序号
    if add_numbering:
        html_text_with_numbers = add_numbering_to_headers(preprocessed_html, numbering_type)
    else:
        html_text_with_numbers = preprocessed_html
    
    # 使用 BeautifulSoup 解析处理后的 HTML
    soup = BeautifulSoup(html_text_with_numbers, 'html.parser')
    
    # 如果没有h1标题，添加传入的title作为文档标题
    if not has_h1_title and title:
        heading = doc.add_heading(title, 0)
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        for run in heading.runs:
            # run.font.color.rgb = RGBColor(0, 0, 128)
            run.font.size = Pt(16)
    
    # 收集所有标题信息
    headers = []
    for tag in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6']):
        level = int(tag.name[1])
        headers.append((tag, level))
    
    # 收集带有strong标签的段落标题
    for p_tag in soup.find_all('p'):
        strong_tags = p_tag.find_all(['strong', 'b'])
        if strong_tags and len(strong_tags) == 1 and strong_tags[0].get_text().strip() == p_tag.get_text().strip():
            text = p_tag.get_text().strip()
            
            def extract_numbering_pattern(text):
                match = re.match(r'^(\d+\.\d+(\.\d+)*)[\s\.]+', text)
                if match:
                    return match.group(1)
                return None
            
            numbering_pattern = extract_numbering_pattern(text)
            level = 3  # 默认三级标题
            
            if numbering_pattern:
                dots_count = numbering_pattern.count('.')
                level = dots_count + 1
                level = max(2, min(level, 6))
            
            headers.append((p_tag, level))
    
    # 创建样式
    def create_styles(doc):
        styles = doc.styles
        if 'List Bullet' not in styles:
            styles.add_style('List Bullet', WD_STYLE_TYPE.PARAGRAPH)
        if 'List Number' not in styles:
            styles.add_style('List Number', WD_STYLE_TYPE.PARAGRAPH)
    
    create_styles(doc)
    
    # 辅助函数定义
    def is_numbered_list_item(text):
        return bool(re.match(r'^\d+\.\s*', text.strip()))
    
    def is_duplicate_numbered_list_item(text):
        return bool(re.match(r'^\d+\.\s+\d+\.\s*', text.strip()))
    
    def process_duplicate_numbered_item(text):
        match = re.match(r'^(\d+\.\s+)(\d+\.\s*)(.*)', text.strip())
        if match:
            first_number, second_number, content = match.groups()
            return first_number, content
        return None, text
    
    # 递归处理HTML元素
    def process_element(element, parent_paragraph=None):
        if element.name is None:  # 文本节点
            text = element.strip()
            if text and parent_paragraph:
                # 检查是否是重复序号的列表项
                if is_duplicate_numbered_list_item(text):
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Pt(18)
                    
                    number, content = process_duplicate_numbered_item(text)
                    if number:
                        run = p.add_run(number)
                        run.bold = True
                        p.add_run(content)
                    else:
                        p.add_run(text)
                # 检查是否是普通数字列表项
                elif is_numbered_list_item(text) and parent_paragraph.style.name != 'List Number':
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Pt(18)
                    
                    match = re.match(r'^(\d+\.\s*)(.*)', text)
                    if match:
                        number, content = match.groups()
                        run = p.add_run(number)
                        run.bold = True
                        p.add_run(content)
                else:
                    parent_paragraph.add_run(text)
            return
        
        # 处理表格
        if element.name == 'table':
            process_table(element, doc)
            return
        
        # 处理图片
        if element.name == 'img':
            process_image(element, doc)
            return
        
        # 处理标题
        if element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
            level = int(element.name[1])
            p = doc.add_paragraph()
            p.style = f'Heading {min(level, 9)}'
            
            run = p.add_run(element.get_text().strip())
            
            font = run.font
            font.size = Pt(16 - level)
            
            if level == 1:
                # font.color.rgb = RGBColor(0, 0, 128)
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            elif level == 2:
                # font.color.rgb = RGBColor(0, 64, 128)
                p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            else:
                p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            
            return
        
        # 处理段落
        if element.name == 'p':
            text = element.get_text().strip()
            
            # 检查段落中是否包含strong标签作为标题
            strong_tags = element.find_all(['strong', 'b'])
            if strong_tags and len(strong_tags) == 1 and strong_tags[0].get_text().strip() == text:
                p = doc.add_paragraph()
                
                level = 3  # 默认值
                for tag, lvl in headers:
                    if tag == element:
                        level = lvl
                        break
                
                p.style = f'Heading {min(level, 9)}'
                run = p.add_run(text)
                
                font = run.font
                font.size = Pt(16 - level)
                
                if level == 1:
                    # font.color.rgb = RGBColor(0, 0, 128)
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                elif level == 2:
                    # font.color.rgb = RGBColor(0, 64, 128)
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                else:
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                
                return
            
            # 处理列表项
            if is_duplicate_numbered_list_item(text):
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Pt(18)
                
                number, content = process_duplicate_numbered_item(text)
                if number:
                    run = p.add_run(number)
                    run.bold = True
                    p.add_run(content)
                else:
                    p.add_run(text)
            elif is_numbered_list_item(text):
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Pt(18)
                
                match = re.match(r'^(\d+\.\s*)(.*)', text)
                if match:
                    number, content = match.groups()
                    run = p.add_run(number)
                    run.bold = True
                    p.add_run(content)
            else:
                p = doc.add_paragraph()
                # 设置首行缩进为2个字符（约等于2个中文字符的宽度）
                p.paragraph_format.first_line_indent = Pt(21)  # 约等于2个中文字符的宽度
                for child in element.children:
                    process_element(child, p)
            return
        
        # 处理格式化标签
        if element.name in ['strong', 'b']:
            if parent_paragraph:
                run = parent_paragraph.add_run(element.get_text().strip())
                run.bold = True
            return
        
        if element.name in ['em', 'i']:
            if parent_paragraph:
                run = parent_paragraph.add_run(element.get_text().strip())
                run.italic = True
            return
        
        if element.name == 'u':
            if parent_paragraph:
                run = parent_paragraph.add_run(element.get_text().strip())
                run.underline = True
            return
        
        # 处理列表
        if element.name == 'ul':
            for li in element.find_all('li', recursive=False):
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Pt(18)
                run = p.add_run('• ')
                run.bold = True
                for child in li.children:
                    process_element(child, p)
            return
        
        if element.name == 'ol':
            for i, li in enumerate(element.find_all('li', recursive=False), 1):
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Pt(18)
                run = p.add_run(f'{i}. ')
                run.bold = True
                for child in li.children:
                    process_element(child, p)
            return
        
        # 处理其他元素
        for child in element.children:
            process_element(child, parent_paragraph)
    
    # 处理HTML主体
    for element in soup.body.children if soup.body else soup.children:
        if isinstance(element, str):
            text = element.strip()
            if text:
                if is_duplicate_numbered_list_item(text):
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Pt(18)
                    number, content = process_duplicate_numbered_item(text)
                    if number:
                        run = p.add_run(number)
                        run.bold = True
                        p.add_run(content)
                    else:
                        p.add_run(text)
                elif is_numbered_list_item(text):
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Pt(18)
                    match = re.match(r'^(\d+\.\s*)(.*)', text)
                    if match:
                        number, content = match.groups()
                        run = p.add_run(number)
                        run.bold = True
                        p.add_run(content)
                else:
                    p = doc.add_paragraph()
                    run = p.add_run(text)
                    font = run.font
                    font.size = Pt(10.5)
                    # 设置首行缩进为2个字符（约等于2个中文字符的宽度）
                    p.paragraph_format.first_line_indent = Pt(21)  # 约等于2个中文字符的宽度
        else:
            process_element(element)
    
    # 修复文档中的标题格式
    fixed_count = fix_document_numbering(doc)
    
    # 如果需要添加版本历史
    if versions:
        doc.add_page_break()
        history_heading = doc.add_heading("版本历史", 1)
        history_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        for v in versions:
            p = doc.add_paragraph()
            p.add_run(f"版本 {v['version']} - {v['created_at']}").bold = True
            if v.get('comment'):
                p.add_run(f" - {v['comment']}")
    
    # 将文档保存到内存中
    docx_bytes = BytesIO()
    doc.save(docx_bytes)
    docx_bytes.seek(0)
    
    return docx_bytes

def html_to_pdf(
    html_content: str, 
    title: str = "Document", 
    author: str = "System",
    versions: Optional[List[dict]] = None,
    add_numbering: bool = True,
    numbering_type: str = "number",
) -> BytesIO:
    """
    将HTML内容转换为PDF格式
    
    Args:
        html_content: HTML格式的文档内容
        title: 文档标题
        author: 文档作者
        versions: 可选的版本历史列表，每个版本是一个包含version, content, comment, created_at的字典
        add_numbering: 是否为标题添加序号，默认为True
        
    Returns:
        BytesIO: 包含PDF文件内容的二进制流
    """
    try:
        import pdfkit
        from io import BytesIO
        import tempfile
        
        # 预处理HTML，合并列表项后的内容
        def preprocess_html(html_text):
            # 使用正则表达式查找列表项模式并合并后续内容
            pattern = r'(\d+\.)\s*\n+\s*([^<>\d\n][^<>\n]*)'
            processed_html = re.sub(pattern, r'\1 \2', html_text)
            
            # 处理HTML中的有序列表
            soup = BeautifulSoup(processed_html, 'html.parser')
            for ol in soup.find_all('ol'):
                for li in ol.find_all('li'):
                    for p in li.find_all('p'):
                        p.replace_with(p.get_text())
            
            return str(soup)
        
        def process_images_for_pdf(soup):
            """处理HTML中的图片元素，确保它们在PDF中正确显示"""
            for img in soup.find_all('img'):
                src = img.get('src', '')
                if src.startswith('data:image'):
                    # base64图片可以直接使用
                    continue
                elif src.startswith(('http://', 'https://')):
                    try:
                        import requests
                        response = requests.get(src)
                        import base64
                        img_data = base64.b64encode(response.content).decode()
                        content_type = response.headers.get('content-type', 'image/jpeg')
                        img['src'] = f'data:{content_type};base64,{img_data}'
                    except Exception as e:
                        logging.error(f"处理网络图片失败: {str(e)}")
                
                # 添加图片标题
                if img.get('alt'):
                    caption_div = soup.new_tag('div')
                    caption_div['class'] = 'image-caption'
                    caption_div.string = img['alt']
                    img.insert_after(caption_div)
        
        # 预处理HTML
        preprocessed_html = preprocess_html(html_content)
        
        # 检查文档是否有h1标题
        soup_check = BeautifulSoup(preprocessed_html, 'html.parser')
        has_h1_title = bool(soup_check.find('h1'))
        
        # 根据add_numbering参数决定是否添加序号
        if add_numbering:
            html_text_with_numbers = add_numbering_to_headers(preprocessed_html, numbering_type)
        else:
            html_text_with_numbers = preprocessed_html
        
        # 如果没有h1标题，添加传入的title作为文档标题
        if not has_h1_title and title:
            title_html = f"<h1 style='text-align:center;color:#000000;font-size:16pt;'>{title}</h1>"
            html_text_with_numbers = title_html + html_text_with_numbers
        
        # 添加版本历史（如果需要）
        if versions:
            versions_html = "<div style='page-break-before: always;'>"
            versions_html += "<h1 style='text-align:center;'>版本历史</h1>"
            for v in versions:
                versions_html += f"<p><strong>版本 {v['version']} - {v['created_at']}</strong>"
                if v.get('comment'):
                    versions_html += f" - {v['comment']}"
                versions_html += "</p>"
            versions_html += "</div>"
            html_text_with_numbers += versions_html
        
        # 添加基本样式
        css_text = """
        body {
            font-family: SimSun, Arial, sans-serif;
            font-size: 10.5pt;
            line-height: 1.5;
            margin: 2cm;
        }
        h1 {
            font-size: 16pt;
            color: #000000;
            text-align: center;
            margin-top: 24pt;
            margin-bottom: 12pt;
        }
        h2 {
            font-size: 14pt;
            color: #000000;
            margin-top: 18pt;
            margin-bottom: 9pt;
            text-align: left;
        }
        h3 {
            font-size: 13pt;
            color: #000000;
            margin-top: 12pt;
            margin-bottom: 6pt;
            text-align: left;
        }
        h4 {
            font-size: 12pt;
            color: #000000;
            margin-top: 12pt;
            margin-bottom: 6pt;
            text-align: left;
        }
        h5, h6 {
            font-size: 11pt;
            color: #000000;
            margin-top: 12pt;
            margin-bottom: 6pt;
            text-align: left;
        }
        p {
            margin-top: 6pt;
            margin-bottom: 6pt;
            text-indent: 2em;  /* 添加首行缩进 */
        }
        ol, ul {
            margin-left: 18pt;
        }
        li {
            margin-bottom: 6pt;
        }
        strong {
            font-weight: bold;
        }
        em {
            font-style: italic;
        }
        .numbered-list {
            margin-left: 18pt;
        }
        .numbered-list-item {
            font-weight: bold;
        }
        .numbered-list-content {
            font-weight: normal;
        }

        /* 表格样式 */
        .doc-table {
            width: 100%;
            border-collapse: collapse;
            margin: 10pt 0;
            table-layout: fixed;
        }
        
        .doc-table, .doc-table th, .doc-table td {
            border: 1px solid black;
            padding: 8pt;
        }
        
        .table-header {
            background-color: #f2f2f2;
            font-weight: bold;
            text-align: center;
        }
        
        .table-cell {
            text-align: center;
            vertical-align: middle;
        }
        
        /* 确保表格边框在 PDF 中显示 */
        table {
            -webkit-print-color-adjust: exact;
            border-collapse: collapse;
            border: 1px solid black;
        }
        
        th, td {
            border: 1px solid black;
            padding: 8pt;
        }

        /* 表格内的段落不缩进 */
        td p, th p {
            text-indent: 0;
            margin: 0;
        }

        /* 图片样式 */
        img {
            max-width: 100%;
            height: auto;
            display: block;
            margin: 10pt auto;
        }
        
        .image-caption {
            text-align: center;
            font-style: italic;
            margin-top: 5pt;
            margin-bottom: 10pt;
        }
        """
        
        # 处理列表项的样式
        soup = BeautifulSoup(html_text_with_numbers, 'html.parser')
        process_images_for_pdf(soup)
        html_text_with_numbers = str(soup)
        
        # 处理数字列表项
        def is_numbered_list_item(text):
            return bool(re.match(r'^\d+\.\s*', text.strip()))
        
        def is_duplicate_numbered_list_item(text):
            return bool(re.match(r'^\d+\.\s+\d+\.\s*', text.strip()))
        
        # 处理段落中的列表项
        for p in soup.find_all('p'):
            text = p.get_text().strip()
            if is_duplicate_numbered_list_item(text):
                match = re.match(r'^(\d+\.\s+)(\d+\.\s*)(.*)', text.strip())
                if match:
                    first_number, second_number, content = match.groups()
                    p['class'] = p.get('class', []) + ['numbered-list']
                    p.clear()
                    number_span = soup.new_tag('span')
                    number_span['class'] = ['numbered-list-item']
                    number_span.string = first_number
                    p.append(number_span)
                    content_span = soup.new_tag('span')
                    content_span['class'] = ['numbered-list-content']
                    content_span.string = content
                    p.append(content_span)
            elif is_numbered_list_item(text):
                match = re.match(r'^(\d+\.\s*)(.*)', text)
                if match:
                    number, content = match.groups()
                    p['class'] = p.get('class', []) + ['numbered-list']
                    p.clear()
                    number_span = soup.new_tag('span')
                    number_span['class'] = ['numbered-list-item']
                    number_span.string = number
                    p.append(number_span)
                    content_span = soup.new_tag('span')
                    content_span['class'] = ['numbered-list-content']
                    content_span.string = content
                    p.append(content_span)
        
        html_text_with_numbers = str(soup)
        
        # 创建完整的HTML文档
        complete_html = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="utf-8">
            <title>{title}</title>
            <style>
                {css_text}
            </style>
        </head>
        <body>
            {html_text_with_numbers}
        </body>
        </html>
        """
        
        # 配置pdfkit选项
        options = {
            'page-size': 'A4',
            'margin-top': '2cm',
            'margin-right': '2cm',
            'margin-bottom': '2cm',
            'margin-left': '2cm',
            'encoding': 'UTF-8',
            'no-outline': None,
            'quiet': ''
        }
        
        # 创建临时文件来保存PDF
        with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as temp_pdf:
            temp_pdf_path = temp_pdf.name
        
        # 生成PDF到临时文件
        pdfkit.from_string(complete_html, temp_pdf_path, options=options)
        
        # 读取临时文件到内存
        with open(temp_pdf_path, 'rb') as f:
            pdf_bytes = BytesIO(f.read())
        
        # 删除临时文件
        import os
        os.unlink(temp_pdf_path)
        
        # 重置文件指针位置
        pdf_bytes.seek(0)
        return pdf_bytes
        
    except Exception as e:
        # 如果pdfkit失败，记录错误并抛出异常
        import logging
        logging.error(f"PDF转换失败: {str(e)}")
        raise Exception(f"PDF转换失败: {str(e)}")